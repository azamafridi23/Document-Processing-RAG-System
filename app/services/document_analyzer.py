import os
import glob
from typing import List, Dict, Any, Optional
import base64
from dotenv import load_dotenv
from pydantic import BaseModel, Field
import json
import shutil
import re

load_dotenv() # Load environment variables from .env file

# --- Project-specific Imports ---
from app.services.aws import AWS

# LangChain Imports
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import PromptTemplate

# DOCX processing
import docx
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# PDF processing
import fitz  # PyMuPDF

# --- Pydantic Model for Structured Output ---

# This is the final, correct format we want.
class ImageDescription(BaseModel):
    """Represents a description of an image with its S3 path."""
    image_path: str = Field(description="The public S3 URL to the extracted image file.")
    description: Optional[str] = Field(description="A brief description of the image content. Should be null if the image is ambiguous.")

# This is a simpler model for the LLM to populate. It doesn't need to know about paths.
class LLMImageDescription(BaseModel):
    """Represents just the description of an image for the LLM to generate."""
    description: Optional[str] = Field(description="A brief, detailed description of the image's content and relevance. Should be null if ambiguous.")

# This is the new model that defines the structure for the LLM's output.
class LLMDocumentAnalysis(BaseModel):
    """Represents the structured analysis from the LLM."""
    document_summary: Optional[str] = Field(description="A concise summary of the entire document. Should be null if the text is empty or ambiguous.")
    image_descriptions: List[LLMImageDescription] = Field(description="A list of descriptions for each image found in the document.")

class DocumentAnalysis(BaseModel):
    """Represents the structured analysis of a document."""
    document_summary: Optional[str] = Field(description="A concise summary of the entire document.")
    # The final output will use the full ImageDescription model.
    image_descriptions: List[ImageDescription] = Field(description="A list of descriptions and S3 paths for each image found in the document.")

# --- Helper to encode image to Base64 ---
def encode_image_to_base64(image_path):
    """Encodes an image file to a Base64 string."""
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except FileNotFoundError:
        print(f"Error: Image file not found at {image_path}")
        return None
    except Exception as e:
        print(f"Error encoding image {image_path}: {e}")
        return None
    
# --- Lightweight Checkers for Image Presence ---

def docx_has_images(file_path: str) -> bool:
    """Quickly checks if a .docx file contains any images without full parsing."""
    try:
        doc = docx.Document(file_path)
        # The presence of inline_shapes is a strong indicator of images.
        if doc.inline_shapes:
            return True
        # Also check for images inside drawing elements in each run
        for p in doc.paragraphs:
            for r in p.runs:
                if 'w:drawing' in r._r.xml:
                    return True
        return False
    except Exception:
        return False

def pdf_has_images(file_path: str) -> bool:
    """Quickly checks if a .pdf file contains any images using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            if page.get_images(full=True):
                return True
        return False
    except Exception:
        return False


# --- DOCX Parsing Logic ---
# (Keeping your original DOCX parsing logic)

def _iter_block_items(parent):
    """
    Yield each paragraph and table child within a parent element in document order.
    The parent can be the document body, a table cell, etc.
    """
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

def _get_image_placeholder(index: int) -> str:
    """Returns a generic placeholder string with an index."""
    return f"[IMAGE_PLACEHOLDER_{index}]"

def parse_docx(file_path: str, file_id: str, image_output_dir: str, aws_manager: AWS, bucket_name: str) -> Dict[str, Any]:
    """
    Parses a .docx file, extracting text, uploading images to S3 with a unique key, and returning S3 URLs and Base64 strings.
    It will only process and upload the first 50 images.
    """
    try:
        doc = docx.Document(file_path)
        full_text_parts = []
        image_details = []
        image_index = 0
        base_filename = os.path.splitext(os.path.basename(file_path))[0]

        os.makedirs(image_output_dir, exist_ok=True)

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                para_text = ""
                for run in block.runs:
                    if 'w:drawing' in run._r.xml:
                        try:
                            for drawing in run._r.xpath("w:drawing"):
                                blip = drawing.find(f".//{qn('a:blip')}")
                                if blip is not None:
                                    r_embed = blip.get(qn('r:embed'))
                                    if r_embed:
                                        if image_index >= 50:
                                            if image_index == 50:  # Log only once
                                                print(f"  - INFO: Document contains more than 50 images. Only processing the first 50.")
                                            # We must still increment the index to keep placeholders unique if ever needed, but skip processing.
                                            full_text_parts.append(_get_image_placeholder(image_index))
                                            image_index += 1
                                            break  # Exit from the drawing loop for this run

                                        related_part = run.part.rels[r_embed].target_part
                                        if related_part.content_type.startswith('image/'):
                                            image_bytes = related_part.blob
                                            image_ext = related_part.content_type.split('/')[-1]
                                            
                                            image_filename = f"{base_filename}_image_{image_index}.{image_ext}"
                                            local_temp_path = os.path.normpath(os.path.join(image_output_dir, image_filename))
                                            with open(local_temp_path, "wb") as f:
                                                f.write(image_bytes)
                                            
                                            # Use file_id to create a unique S3 key
                                            s3_key = f"images/{file_id}_{image_filename}"
                                            aws_manager.upload_file_to_s3(local_temp_path, bucket_name, s3_key)
                                            s3_url = aws_manager.generate_object_url(bucket_name, s3_key)
                                            base64_image = encode_image_to_base64(local_temp_path)
                                            
                                            os.remove(local_temp_path)

                                            if s3_url and base64_image:
                                                image_details.append({"s3_url": s3_url, "base64": base64_image})
                                                full_text_parts.append(_get_image_placeholder(image_index))
                                                image_index += 1
                                            break
                        except Exception as img_e:
                            print(f"  - Warning: Failed to process an image in {file_path}. Error: {img_e}")
                    else:
                        para_text += run.text
                full_text_parts.append(para_text)

            elif isinstance(block, Table):
                for row in block.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    full_text_parts.append(row_text)
            
            full_text_parts.append("\\n")

        return {"text_content": "".join(full_text_parts), "image_details": image_details}

    except Exception as e:
        print(f"Error processing DOCX file {file_path}: {e}")
        return {"text_content": "", "image_details": []}


# --- PDF Parsing Logic ---

def parse_pdf(file_path: str, file_id: str, image_output_dir: str, aws_manager: AWS, bucket_name: str) -> Dict[str, Any]:
    """
    Parses a .pdf file, extracting text, uploading images to S3 with a unique key, and returning S3 URLs and Base64 strings.
    It will only process and upload the first 50 images.
    """
    try:
        doc = fitz.open(file_path)
        full_text_content_parts = []
        image_details = []
        base_filename = os.path.splitext(os.path.basename(file_path))[0]

        os.makedirs(image_output_dir, exist_ok=True)

        # This index must be document-wide, not per-page, to create unique placeholders.
        global_image_index = 0

        for page_num, page in enumerate(doc):
            full_text_content_parts.append(f"--- Page {page_num + 1} ---\\n")
            page_data = page.get_text("dict", sort=True)
            blocks = page_data["blocks"]
            
            # This index is only for the local filename to ensure uniqueness per page.
            page_image_index = 0

            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block["lines"]:
                        for span in line["spans"]:
                            full_text_content_parts.append(span["text"])
                        full_text_content_parts.append("\\n")
                elif block["type"] == 1:  # Image block
                    try:
                        if global_image_index >= 50:
                            if global_image_index == 50:  # Log only once
                                print(f"  - INFO: Document contains more than 50 images. Only processing the first 50.")
                            # We must still increment the index to keep placeholders unique, but skip processing.
                            full_text_content_parts.append(_get_image_placeholder(global_image_index))
                            global_image_index += 1
                            page_image_index += 1
                            continue  # Skip to the next block

                        image_bytes = block["image"]
                        image_ext = block["ext"]
                        
                        image_filename = f"{base_filename}_page_{page_num+1}_image_{page_image_index}.{image_ext}"
                        local_temp_path = os.path.normpath(os.path.join(image_output_dir, image_filename))
                        with open(local_temp_path, "wb") as f:
                            f.write(image_bytes)

                        # Use file_id to create a unique S3 key
                        s3_key = f"images/{file_id}_{image_filename}"
                        aws_manager.upload_file_to_s3(local_temp_path, bucket_name, s3_key)
                        s3_url = aws_manager.generate_object_url(bucket_name, s3_key)
                        base64_image = encode_image_to_base64(local_temp_path)
                        
                        os.remove(local_temp_path)

                        if s3_url and base64_image:
                            image_details.append({"s3_url": s3_url, "base64": base64_image})
                            # Use the corrected global_image_index for the placeholder
                            full_text_content_parts.append(_get_image_placeholder(global_image_index))
                            global_image_index += 1
                        
                        page_image_index += 1

                    except Exception as e:
                        print(f"Warning: Could not extract/upload an image from {file_path} on page {page_num+1}: {e}")

        return {"text_content": "".join(full_text_content_parts), "image_details": image_details}
    except Exception as e:
        print(f"Error processing PDF file {file_path}: {e}")
        return {"text_content": "", "image_details": []}


# --- Simple Text Extraction Functions ---

def parse_pdf_text(file_path: str) -> str:
    """Extracts and returns all text content from a PDF file."""
    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        return "\\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return ""

def parse_docx_text(file_path: str) -> str:
    """Extracts and returns all text content from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        # You might also want to extract text from tables if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return ""


# --- Main Document Processing and OpenAI Integration Function (using LangChain) ---

def analyze_document_with_openai_langchain_structured(
    file_path: str,
    file_id: str,
    user_initial_prompt: str,
    model: str = "gpt-4o-mini",
    image_output_dir: str = "extracted_document_images" # This is now a temporary directory
) -> Dict[str, Any]:
    """
    Analyzes a DOCX or PDF file, uploads images to S3, and returns structured output.
    """
    # Create a temporary directory for this specific run to avoid conflicts
    temp_image_dir = os.path.join(image_output_dir, f"temp_{os.path.basename(file_path)}_{os.getpid()}")
    os.makedirs(temp_image_dir, exist_ok=True)
    
    parsed_data = {"text_content": "", "image_details": []}
    aws_manager = AWS()
    s3_bucket_name = os.getenv("S3_BUCKET_NAME")
    if not s3_bucket_name:
        raise ValueError("S3_BUCKET_NAME environment variable is not set.")

    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")

        if file_path.lower().endswith(".pdf"):
            parsed_data = parse_pdf(file_path, file_id, temp_image_dir, aws_manager, s3_bucket_name)
        elif file_path.lower().endswith(".docx"):
            parsed_data = parse_docx(file_path, file_id, temp_image_dir, aws_manager, s3_bucket_name)
        else:
            raise ValueError("Unsupported file type. Only .docx and .pdf are supported.")

        if not parsed_data["text_content"] and not parsed_data["image_details"]:
            raise ValueError(f"Could not extract any content (text or images) from {file_path}. Document might be empty or corrupted.")

        # Construct the multimodal content list for LangChain's HumanMessage
        langchain_content = []
        text_with_placeholders = parsed_data["text_content"]
        image_details_list = parsed_data["image_details"]

        # This regex was incorrect, it was looking for a literal backslash.
        # The new pattern correctly finds "[IMAGE_PLACEHOLDER_...]"
        pattern = r"\[IMAGE_PLACEHOLDER_(\d+)\]"
        last_index = 0

        for match in re.finditer(pattern, text_with_placeholders):
            start, end = match.span()
            # Add text before the placeholder
            pre_text = text_with_placeholders[last_index:start].strip()
            if pre_text:
                langchain_content.append({"type": "text", "text": pre_text})
            
            # Get image index from placeholder (group 1 now contains the digits)
            image_index = int(match.group(1))
            
            if image_index < len(image_details_list):
                details = image_details_list[image_index]
                base64_image = details.get('base64')
                s3_url = details.get('s3_url')

                if base64_image and s3_url:
                    # Add the image data (Base64) for the API
                    image_ext = os.path.splitext(s3_url)[1].lstrip('.')
                    image_type = f"image/{image_ext}" if image_ext in ["png", "jpeg", "jpg", "gif", "webp"] else "image/jpeg"
                    langchain_content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{image_type};base64,{base64_image}",
                            "detail": "high"
                        }
                    })

                    # Add the S3 URL for the LLM to see in the text context
                    langchain_content.append({"type": "text", "text": f"Image is available at: {s3_url}"})
            
            last_index = end

        # Add any remaining text after the last placeholder
        final_text = text_with_placeholders[last_index:].strip()
        if final_text:
            langchain_content.append({"type": "text", "text": final_text})
        
        # --- LangChain Output Parsing Setup ---
        # The parser now targets the new, simpler model for the LLM's output, fixing the error.
        parser = PydanticOutputParser(pydantic_object=LLMDocumentAnalysis)

        # Create the prompt for the LLM
        structured_prompt_template = PromptTemplate(
            template="""You are an expert document analyst. Your task is to analyze the provided document content and images.
            The user's request is: {user_initial_prompt}

The document content may have text and image placeholders. You will be given the images corresponding to these placeholders.
            
**Analysis Rules and Handling Ambiguity:**
1.  Provide a concise summary of the document's text. If the document's text is empty, nonsensical, or too ambiguous to summarize, the value for the `document_summary` field MUST be `null`.
2.  Provide a detailed description for EACH image, in the order they are presented. If an individual image is blank, corrupted, or too ambiguous to describe, its `description` field MUST be `null`.
3.  If there are no images provided, `image_descriptions` must be an empty list `[]`.
            
Provide your final answer in the specified JSON format, following the schema precisely. Do not include placeholders in your output.

            {format_instructions}
            """,
            input_variables=["user_initial_prompt"],
            partial_variables={"format_instructions": parser.get_format_instructions()},
        )
        
        # Combine the structured prompt with the multimodal document content
        final_prompt_content = langchain_content + [{"type": "text", "text": structured_prompt_template.format(user_initial_prompt=user_initial_prompt)}]

        # --- TEMPORARY DEBUGGING STEP: Verify final payload ---
        print("\n--- Verifying final content being sent to LLM ---")
        payload_has_image = False
        for i, item in enumerate(final_prompt_content):
            if item["type"] == "text":
                # Print first 200 chars of text parts for context
                print(f"  - Part {i} (text): {item['text'][:200].strip()}...")
            elif item["type"] == "image_url":
                # Check if the image_url part is correctly formatted and has data
                url_data = item.get("image_url", {}).get("url", "")
                print(f"  - Part {i} (image_url): Image is present. Base64 data length: {len(url_data)}")
                payload_has_image = True
        
        if not payload_has_image:
            print("  - ⚠️ WARNING: No image data was found in the final payload sent to the LLM.")
        print("-------------------------------------------------\n")

        # Initialize LangChain's ChatOpenAI model
        llm = ChatOpenAI(model=model, temperature=0.1) # Slightly higher temp, but still focused on structure

        # Create a HumanMessage with the multimodal content
        messages = [
            HumanMessage(content=final_prompt_content)
        ]

        # Invoke the model and parse the output
        llm_response = llm.invoke(messages)

        # --- TEMPORARY DEBUGGING STEP ---
        print("--- RAW LLM RESPONSE ---")
        print(llm_response.content)
        print("------------------------")
        
        # Try to parse the output using the Pydantic parser
        parsed_output: LLMDocumentAnalysis = parser.parse(llm_response.content)

        # --- Reconstruct the final data structure ---
        # Now, combine the LLM's descriptions with the S3 URLs we already have.
        final_image_descriptions = []
        if parsed_output.image_descriptions:
            for i, desc_item in enumerate(parsed_output.image_descriptions):
                if i < len(image_details_list):
                    s3_url = image_details_list[i]["s3_url"]
                    final_image_descriptions.append({
                        "image_path": s3_url,
                        "description": desc_item.description
                    })

        # Return the final, correctly formatted dictionary including complete text
        return {
            "document_summary": parsed_output.document_summary,
            "image_descriptions": final_image_descriptions,
            "complete_document_text": parsed_data["text_content"]
        }

    except Exception as e:
        # Catch any errors during parsing or API call
        print(f"An error occurred during document analysis: {e}")
        # Try to return the raw content if Pydantic parsing failed, for debugging
        if 'llm_response' in locals():
            print(f"Raw LLM Response Content (could not parse): {llm_response.content}")
            return {
                "error": str(e), 
                "raw_llm_response": llm_response.content,
                "complete_document_text": parsed_data.get("text_content", "") if 'parsed_data' in locals() else ""
            }
        return {
            "error": str(e),
            "complete_document_text": parsed_data.get("text_content", "") if 'parsed_data' in locals() else ""
        }
    finally:
        # Clean up the temporary image directory for this run
        if os.path.exists(temp_image_dir):
            shutil.rmtree(temp_image_dir)
            print(f"Cleaned up temporary image directory: {temp_image_dir}")


if __name__ == '__main__':
    print("--- OpenAI Document Analysis Tool (with S3 Upload) ---")
    print("This tool extracts content, analyzes it with OpenAI, and provides structured JSON output with actual image paths.")
    print("Extracted images will remain in the specified folder.")
    
    # --- User Input Section ---
    # You can directly paste the full path to your file here.
    # Examples:
    # my_document_path = "/home/youruser/Documents/my_report.docx"
    # my_document_path = "C:\\Users\\YourUser\\Desktop\\image_document.pdf"

    # IMPORTANT: REPLACE THIS WITH YOUR ACTUAL FILE PATH
    my_document_path = r"/path/to/project/app/data_ingestion/google_drive_files/Temp1/SOP Entering Amazon Order in QB.docx"
    # my_document_path = r"path/to/your/pdf_document.pdf" # Uncomment and replace for PDF

    # The directory where extracted images will be saved. It will not be deleted.
    # You can change this to any path you prefer, e.g., "my_extracted_images_folder"
    persistent_image_output_dir = "extracted_document_images" 

    # You can modify the prompt as needed. This prompt will be combined with the structured output instructions.
    user_prompt = "Summarize this document in detail and provide a detailed description of each image. Explain its relevance to the document content. For each image description, clearly state the full file path where the image is saved on the system."

    if not my_document_path:
        print("No file path provided. Exiting.")
    elif not os.path.exists(my_document_path):
        print(f"Error: The file path '{my_document_path}' does not exist.")
    else:
        print(f"\n--- Analyzing file: {my_document_path} ---")
        
        # Call the new structured analysis function
        analysis_result = analyze_document_with_openai_langchain_structured(
            my_document_path, 
            "some_file_id", # Replace with a real file_id for standalone testing
            user_prompt, 
            image_output_dir=persistent_image_output_dir # Use the persistent directory
        )
        
        print("\n--- OpenAI Structured Analysis Result ---")
        if isinstance(analysis_result, dict) and "error" in analysis_result:
            print(f"Analysis failed: {analysis_result['error']}")
            if 'raw_llm_response' in analysis_result:
                print(f"Raw LLM Response: {analysis_result['raw_llm_response']}")
        else:
            print(json.dumps(analysis_result, indent=2)) # Print the structured dict beautifully
        print("\n--- Analysis Complete ---")